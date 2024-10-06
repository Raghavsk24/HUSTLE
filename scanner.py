
import docx
from PyPDF2 import PdfReader
import PyPDF2

# Dictionary for side hustles
side_hustles = {
    "Scriptwriter": ["Creative writing", "Screenwriting", "Story development", "Dialogue creation", "Narrative structure", "English", "Content", "Media"],
    "Prompt Engineer": ["Natural language processing", "AI", "Model optimization", "GPT-3/4", "Data analysis", "Computer Science", "Programming", "Technical"],
    "Medium Content Writer": ["Blog writing", "SEO optimization", "Content strategy", "Article development", "Copywriting", "English", "Content", "Media"],
    "Graphic Designer": ["Adobe Creative Suite", "Branding", "Visual communication", "Typography", "Layout design", "Art", "Creativity", "Technical"],
    "Animator": ["2D/3D animation", "Motion graphics", "Storyboarding", "After Effects", "Character design", "Programming", "Computer science", "Technical"],
    "Photographer": ["Photo editing", "Lightroom", "Event photography", "Composition", "Lighting techniques", "Entrepreneurship", "Content", "Media"],
    "3D Printing Specialist": ["CAD", "Additive manufacturing", "3D modeling", "Prototyping", "AutoCAD", "Graphic Design", "Technical", "Math"],
    "Virtual Bookkeeper": ["QuickBooks", "Accounts payable/receivable", "Financial reporting", "Reconciliation", "Data entry", "Analysis", "Management", "Technical"],
    "Podcast Host": ["Public speaking", "Audio editing", "Content creation", "Interviewing skills", "Audience engagement", "Communication", "Content", "Media"],
    "Online Tutor": ["Curriculum development", "Lesson planning", "Subject expertise", "Virtual teaching platforms", "Student engagement", "STEM", "Honor Roll", "High School"],
    "Voiceover Artist": ["Voice modulation", "Audio recording", "Diction", "Script interpretation", "Studio equipment", "Content", "Media", "Editing"],
    "Web Developer": ["HTML/CSS/JavaScript", "Full-stack development", "Responsive design", "API integration", "Front-end/back-end development", "Bootstrap", "VSCode", "Django"],
    "Babysitter": ["Childcare", "CPR certification", "Time management", "Conflict resolution", "Activity planning", "Customer Service", "Organization", "Adaptability"],
    "Content Editor": ["Proofreading", "Copy editing", "Grammar and syntax", "Style guides", "Attention to detail", "Content", "Media", "Writing"],
    "Video Editor": ["Final Cut Pro", "Storytelling", "Color grading", "Video transitions", "Timeline management", "Software", "Technical", "Media"],
    "Dog Walker": ["Pet care", "Time management", "Route planning", "Animal behavior", "Reliability", "Organization", "Customer Service", "Adaptability"]
}

# Dictionary for side hustle scores
side_hustle_scores = {side_hustle : 0 for side_hustle in side_hustles }

# Returns text if uploaded file extension is.docx
def process_docx(file_path):
    try:
        doc = docx.Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    except Exception as e:
        print(f"Error reading DOCX resume: {e}")

# Returns text if uploaded file extension is .txt
def process_txt(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except Exception as e:
        print(f"Error reading TXT resume: {e}")
        return ""

# Returns text if uploaded file extension is .pdf
def process_pdf(file_path):
    try:
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
        return text
    except Exception as e:
        print(f"Error reading PDF resume: {e}")
        return ""

def process_resume(file_path):
    # Initalize side hustle scores
    side_hustle_scores = {side_hustle: 0 for side_hustle in side_hustles}  \

    try:
        # Process .docx files
        if file_path.endswith('.docx'):
            doc = docx.Document(file_path)
            for paragraph in doc.paragraphs:
                line = paragraph.text.strip()
                print(f"Processing line from DOCX: {line}") # prints text onto terminal 
                for side_hustle, keywords in side_hustles.items():
                    for keyword in keywords:
                        if keyword.lower() in line.lower():
                            side_hustle_scores[side_hustle] += 1
        
        # Process .txt files
        elif file_path.endswith('.txt'):
            with open(file_path, 'r') as file:
                lines = file.readlines()
                for line in lines:
                    line = line.strip()
                    print(f"Processing line from TXT: {line}")  # prints text onto terminal 
                    for side_hustle, keywords in side_hustles.items():
                        for keyword in keywords:
                            if keyword.lower() in line.lower():
                                side_hustle_scores[side_hustle] += 1

        # Process .pdf files
        elif file_path.endswith('.pdf'):
            with open(file_path, 'rb') as file: 
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        for line in text.split('\n'):
                            line = line.strip()
                            print(f"Processing line from PDF: {line}")  # prints text onto terminal 
                            for side_hustle, keywords in side_hustles.items():
                                for keyword in keywords:
                                    if keyword.lower() in line.lower():
                                        side_hustle_scores[side_hustle] += 1

    except Exception as e:
        print(f"Error processing resume: {e}")

    
    # Print highest side hustles
    if side_hustle_scores:
        highest_hustle = max(side_hustle_scores, key=side_hustle_scores.get)
        highest_score = side_hustle_scores[highest_hustle]
        if highest_score > 0:
            print(f"Highest Hustle: {highest_hustle}, Score: {highest_score}")  
            return highest_hustle, highest_score
    return None, None  

